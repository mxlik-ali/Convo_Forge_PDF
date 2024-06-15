import os
from PyPDF2 import PdfReader
from docx import Document
from openpyxl import load_workbook
import pandas as pd
import html

class DocumentClassifier:
    """
    A class to classify and extract content from various document types.
    
    This class supports PDF, DOCX, XLSX, and CSV file formats.
    
    Attributes:
        document (str): The path to the document to be processed.
    """
    
    def __init__(self, document: str) -> None:
        """
        Initialize the DocumentClassifier with the path to the document.
        
        Parameters:
            document (str): The path to the document.
        """
        self.document = document

    def pdf_get_content(self) -> str:
        """
        Extract text from a PDF document.
        
        This method uses the PyPDF2 library to read the text content from each page of the PDF.
        
        Returns:
            str: The extracted text content as a single string, with each page separated by two newlines.
        """
        text = []

        try:
            pdf_reader = PdfReader(self.document)
            num_pages = len(pdf_reader.pages)
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text.append(f"Page {page_num + 1}:\n{page_text}\n")
        except Exception as e:
            print(f"Error reading PDF file: {e}")
        
        return "\n\n".join(text)

    def read_docs(self) -> str:
        """
        Extract text and tables from a DOCX document.
        
        This method uses the python-docx library to read text from paragraphs and cells of tables 
        in a DOCX document. The paragraphs and tables are combined into a single string.
        
        Returns:
            str: The extracted text content, with paragraphs and tables formatted as plain text.
        """
        try:
            doc = Document(self.document)
            paragraphs = [para.text for para in doc.paragraphs]
            tables = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    tables.append(" | ".join(row_text))
                tables.append("\n")  # Add a newline after each table

            return "\n\n".join(paragraphs + tables)
        except Exception as e:
            print(f"Error reading DOCX file: {e}")
            return ""

    def read_xlsx(self) -> str:
        """
        Extract data from an XLSX document.
        
        This method uses the openpyxl library to read data from each sheet in an XLSX file. The data 
        from each sheet is combined into a single string.
        
        Returns:
            str: The extracted data, with each sheet and row formatted as plain text.
        """
        try:
            workbook = load_workbook(filename=self.document)
            sheets_data = []
            for sheetname in workbook.sheetnames:
                sheet = workbook[sheetname]
                sheet_data = [sheetname]
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(map(str, row))
                    sheet_data.append(row_text)
                sheets_data.append("\n".join(sheet_data))

            return "\n\n".join(sheets_data)
        except Exception as e:
            print(f"Error reading XLSX file: {e}")
            return ""

    def read_csv(self) -> str:
        """
        Extract data from a CSV document.
        
        This method uses the pandas library to read data from a CSV file into a DataFrame and then 
        convert it to a plain text string.
        
        Returns:
            str: The extracted data as plain text.
        """
        try:
            data = pd.read_csv(self.document)
            return data.to_string(index=False)
        except Exception as e:
            print(f"Error reading CSV file: {e}")
            return ""

    def process_file(self) -> str:
        """
        Process the document based on its file extension.
        
        This method determines the file type by examining the file extension and calls the 
        appropriate method to extract content from the document. If the file type is not supported, 
        it raises a ValueError.
        
        Returns:
            str: The extracted content from the document as a plain text string.
        
        Raises:
            ValueError: If the file type is unsupported.
        """
        ext = os.path.splitext(self.document)[1].lower()
        if ext == '.pdf':
            return self.pdf_get_content()
        elif ext == '.csv':
            return self.read_csv()
        elif ext == '.xlsx':
            return self.read_xlsx()
        elif ext == '.docx':
            return self.read_docs()
        else:
            raise ValueError(f"Unsupported file type: {ext}")









# Example Usage
# process = DocumentClassifier(document="C:\Convo_Forge_PDF\data\docs_2\PDF-Guide-Node-Andrew-Mead-v3.docx")
# print(process)

# processed_output = process.process_file()
# with open('example.txt','w')as f:
#     f.write(str(processed_output))